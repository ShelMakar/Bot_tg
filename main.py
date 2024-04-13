import sqlite3
import telebot
from telebot import types
from docx import Document
import math
import os
import time


base_path= 'C:\моя папочка\phyton\projects\BOT-stroitel'
TOKEN = '6745250256:AAH69sfCYOmZecFGtj3EqxO1xj8FM4pTFj0'
bot = telebot.TeleBot(TOKEN)
user_data = {}
# Константы
a = 0.00035
b = 1.4
mp = 1
alpha_int = 8.7
alpha_ext = 23
r = 0.99
#Теплопроводность слоя X для каждого варианта
conductivity_x = {
    'Вариант 1': 0.044,  # Вариант 1
    'Вариант 2': 0.044,  # Вариант 2
    'Вариант 3': 0.042,  # Вариант 3
    'Вариант 4': 0.042  # Вариант 4
}
def fetch_city_data(city_name):
    conn = sqlite3.connect('city.db')
    cursor = conn.cursor()
    query = "SELECT duration, average_temperature FROM heating_periods WHERE city = ?"
    cursor.execute(query, (city_name,))
    data = cursor.fetchone()
    conn.close()
    return data

def calculate_variant(wall_type, wall_material):
    variants = {
        '1.1_2.1': 'Вариант 1',
        '1.1_2.2': 'Вариант 2',
        '1.2_2.1': 'Вариант 3',
        '1.2_2.2': 'Вариант 4'
    }
    key = f"{wall_type}_{wall_material}"
    return variants.get(key, 'Неизвестный вариант')

def calculate_R_0_usl(W, R_0_norm, variant):
    # Константы
    alpha_int = 8.7  # Внутренний коэффициент теплопередачи
    alpha_ext = 23  # Внешний коэффициент теплопередачи

    # Словарь с данными слоев для каждого варианта
    variants = {
        'Вариант 1': [(2, 0.8), (4, 0.8), (W, 0.81)],  # Вариант 1
        'Вариант 2': [(2, 0.8), (4, 0.8), (W, 0.5)],  # Вариант 2
        'Вариант 3': [(8, 0.22), (W, 0.81)],  # Вариант 3
        'Вариант 4': [(8, 0.22), (W, 0.5)]  # Вариант 4
    }

    # Выбираем набор данных слоев на основе варианта
    layers = variants.get(variant, [])
    a_x = conductivity_x.get(variant)

    # Расчет условного сопротивления теплопередачи R_0 усл без учета слоя X
    R_0_usl = 1 / alpha_int
    for thickness, conductivity in layers:
        thickness = float(thickness)
        if thickness is not None:
            R_0_usl += thickness * 0.001 / conductivity
    R_0_usl += 1 / alpha_ext
    print (R_0_norm - R_0_usl)
    print(a_x)
    # Расчет толщины слоя X
    x = (R_0_norm - R_0_usl) * a_x
    print(x)
    x_rounded = math.ceil(x * 100) * 10

    return R_0_usl, x_rounded,a_x

#Функция добавления переменных в шаблон
def replace_placeholders_in_docx(template_path, output_path, calculations):
    doc = Document(template_path)

    for p in doc.paragraphs:
        full_text = p.text
        for key, value in calculations.items():
            placeholder = f"{{{key}}}"  # Формат плейсхолдера
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(round(value , 2)))
        clear_paragraph_runs(p)
        run = p.add_run(full_text)
        run.font.name = "ISOCPEUR"  # Установка шрифта для всего параграфа

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = p.text
                    for key, value in calculations.items():
                        placeholder = f"{{{key}}}"
                        if placeholder in full_text:
                            full_text = full_text.replace(placeholder, str(value))
                    clear_paragraph_runs(p)
                    run = p.add_run(full_text)
                    run.font.name = "ISOCPEUR"

    doc.save(output_path)
def clear_paragraph_runs(paragraph):
    """
    Очищает все runs в параграфе, чтобы добавить обновленный текст в новый run.
    """
    for run in paragraph.runs:
        run.clear()
def thermal_technical_calculation(message,city_name, variant, W):
    # Запрос данных у пользователя

    # Получение данных о городе


    Z_ot, t_ot = fetch_city_data(city_name)

    # Расчет градусо-суток отопительного периода и R_o_mp
    GСОП = (20 - t_ot) * Z_ot
    R_o_mp = a * GСОП + b
    # Расчет R_o_norm
    R_o_norm = R_o_mp * mp
    # Определение варианта стены и расчет R_0_usl и x
    R_0_usl, x_rounded,a_x = calculate_R_0_usl(W, R_o_norm, variant)
    # Вывод результатов
    # Запуск расчетов и сбор данных в словарь

    calculations = {
        "1": city_name,
        "2": GСОП,
        "3": Z_ot,
        "4": t_ot,
        "5": R_o_mp,
        "6": R_o_norm,
        "7": a_x,
        "8": R_0_usl,
        "9": x_rounded,
        "10": W
    }
    for key, value in calculations.items():
        print(f"{key}: {value}")

    tempvar = {
        'Вариант 1': "1",  # Вариант 1
        'Вариант 2': "2",  # Вариант 2
        'Вариант 3': "3",  # Вариант 3
        'Вариант 4': "4"  # Вариант 4
    }
    templvar = tempvar.get(variant)

    # Путь к вашему шаблону DOCX и путь к новому файлу
    template_path = f'Template{templvar}.docx'  # Укажите здесь путь к вашему шаблону DOCX
    output_path = 'Теплотехнический расчёт.docx'  # Имя нового файла с результатами

    # Замена плейсхолдеров в документе и сохранение нового файла
    replace_placeholders_in_docx(template_path, output_path, calculations)
    chat_id = message.chat.id

    # Отправка файла пользователю
    try:
        # Попытка выполнить действие, которое может вызвать ошибку
        with open(output_path, 'rb') as doc:
            bot.send_document(chat_id, doc)
        time.sleep(1)  # Пауза на 1 секунду перед удалением файла
        os.remove(output_path)
    except Exception as e:
        print(f"Произошла ошибка: {e}")




def main_menu(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    itembtn1 = types.KeyboardButton('Рабочая документация')
    itembtn2 = types.KeyboardButton('Исполнительная документация')
    markup.add(itembtn1, itembtn2)
    bot.send_message(message.chat.id, "Выбери тип документации:", reply_markup=markup)

# Изменение функции send_welcome для вызова main_menu
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    bot.send_message(message.chat.id, "Привет! Я помогу тебе с документацией.")
    main_menu(message)  # Вызов основного меню

@bot.message_handler(func=lambda message: message.text == "Назад")
def handle_back(message):
    main_menu(message)  # Возврат в основное меню

@bot.message_handler(func=lambda message: True)
def handle_message(message):
    if message.text == 'Рабочая документация':
        work_doc(message)
    elif message.text == 'Исполнительная документация':
        isp_doc(message)
    elif message.text == 'Инженерные сети':
        send_inzh(message)
    elif message.text == 'Архитектура':
        bot.send_message(message.chat.id, "Документ находится на стадии разработки")
    elif message.text == 'Система канализаций' or 'Система вентиляций' or 'Система водоснабжения':
        send_doc(message)
    elif message.text == 'Теплотехнический расчет' or 'Расчет сечения воздуховодов':
        process_calculation_step(message)

def isp_doc(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    itembtn1 = types.KeyboardButton('Инженерные сети')
    itembtn2 = types.KeyboardButton('Архитектура')
    itembtn_back = types.KeyboardButton('Назад')
    markup.add(itembtn1, itembtn2, itembtn_back)
    bot.send_message(message.chat.id, "Выбери необходимую документацию:", reply_markup=markup)



def send_inzh(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    itembtn1 = types.KeyboardButton('Система канализаций')
    itembtn2 = types.KeyboardButton('Система вентиляций')
    itembtn3 = types.KeyboardButton('Система водоснабжения')
    itembtn_back = types.KeyboardButton('Назад')
    markup.add(itembtn1, itembtn2, itembtn3, itembtn_back)
    bot.send_message(message.chat.id, "Выбери необходимую систему", reply_markup=markup)

def send_doc(message):
    chat_id = message.chat.id
    voda_path =f'{base_path}\Вода'
    if message.text == 'Система канализаций':
        bot.send_document(chat_id, open(f'{voda_path}/Паспорт трубы из сшитого полителена.pdf', 'rb'))
        bot.send_document(chat_id, open(f'{voda_path}/Паспорт шаровый кран VALTEC.pdf', 'rb'))
        bot.send_document(chat_id, open(f'{voda_path}/Сертификат соответсвия трубы из сшитого полиэтелена.pdf', 'rb'))

    elif message.text == 'Система вентиляций':
        bot.send_message(chat_id, "Документ находится на стадии разработки")
    elif message.text == 'Система водоснабжения':
        # Путь к папке "Канаха"
        kanaha_path = f'{base_path}\Канаха'
        bot.send_document(chat_id, open(f'{kanaha_path}/Паспорт трубы PP-R.pdf', 'rb'))
        bot.send_document(chat_id, open(f'{kanaha_path}/Паспорт фасонина PP-R.pdf', 'rb'))
        bot.send_document(chat_id, open(f'{kanaha_path}/Сертификат соответствия трубы PP-R.pdf', 'rb'))

def work_doc(message):
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    itembtn1 = types.KeyboardButton('Теплотехнический расчет')
    itembtn2 = types.KeyboardButton('Расчет сечения воздуховодов')
    itembtn_back = types.KeyboardButton('Назад')
    markup.add(itembtn1, itembtn2, itembtn_back)
    msg = bot.send_message(message.chat.id, "Выбери необходимый вид расчета.", reply_markup=markup)
    bot.register_next_step_handler(msg, process_calculation_step)

def process_calculation_step(message):
    if message.text == 'Теплотехнический расчет':
        msg = bot.send_message(message.chat.id, 'Введите название вашего города:')
        bot.register_next_step_handler(msg, process_city_step)
    elif message.text == 'Расчет сечения воздуховодов':
        msg = bot.send_message(message.chat.id, 'Введи площадь помещения:')
        bot.register_next_step_handler(msg, s_vozdyh)

def s_vozdyh(message):
    global s
    s = float(message.text)
    msg = bot.send_message(message.chat.id, 'Введи высоту помещения:')
    bot.register_next_step_handler(msg, h_vozdyh)

def h_vozdyh(message):
    global h
    h = float(message.text)
    msg = bot.send_message(message.chat.id, 'Введи кол-во диффузоров:')
    bot.register_next_step_handler(msg, n_vozdyh)

def n_vozdyh(message):
    global n, s, h
    n = float(message.text)
    msg = bot.send_message(message.chat.id, 'Введи частоту обновления воздуха в помещении в час:')
    bot.register_next_step_handler(msg, chastota_vozdyh)

def chastota_vozdyh(message):
    global chast
    chast = float(message.text)
    v = s * h * chast / n
    bot.send_message(message.chat.id, f'Объем воздуха, проходящего через диффузор в час: {v} м3')
    x = (v / 1.5 / 3600)
    r = (x / 3.14) ** 0.5
    rounded_r = [100, 125, 140, 160, 180, 200, 250,
                 280, 315, 355, 400, 450, 500, 560,
                 630, 710, 800, 900, 1000, 1120, 1250]  # здесь необходимые числа для округления
    rounded_number = min(rounded_r, key=lambda y: abs(y - r))
    bot.send_message(message.chat.id, f'Сечение круголого воздуховода: {rounded_number}')
    a = round(((x / 3) ** 0.5) * 2, -2) // 2
    if a == 0:
        a = 50
    b = a * 3
    bot.send_message(message.chat.id, f'Сечение прямоугольного воздуховода:{b}x{a}')


def process_city_step(message):
    city_name = message.text
    data = fetch_city_data(city_name)
    if data:
        chat_id = message.chat.id
        if chat_id not in user_data:
            user_data[chat_id] = {}
        user_data[chat_id]['city_name'] = city_name  # Сохраняем название города
        markup = types.InlineKeyboardMarkup()
        itembtn1 = types.InlineKeyboardButton('Двухслойная', callback_data='wall_1.1')
        itembtn2 = types.InlineKeyboardButton('Трехслойная', callback_data='wall_1.2')
        markup.add(itembtn1, itembtn2)
        bot.send_message(message.chat.id, "Выберите тип стены:", reply_markup=markup)
    else:
        msg = bot.send_message(message.chat.id, 'Данные по данному городу не найдены. Пожалуйста, введите название города еще раз:')
        bot.register_next_step_handler(msg, process_city_step)



@bot.callback_query_handler(func=lambda call: True)  # Обрабатываем все callback-запросы
def callback_query(call):
    chat_id = call.message.chat.id

    if chat_id not in user_data:
        user_data[chat_id] = {}

    if call.data.startswith('wall_'):
        user_data[chat_id]['wall_type'] = call.data.split('_')[1]
        # Логика для выбора материала стены
        markup = types.InlineKeyboardMarkup()
        itembtn1 = types.InlineKeyboardButton('Кирпич', callback_data='material_2.1')
        itembtn2 = types.InlineKeyboardButton('Газобетон/пеноблок', callback_data='material_2.2')
        markup.add(itembtn1, itembtn2)
        bot.send_message(chat_id, "Выберите материал стены:", reply_markup=markup)
    elif call.data.startswith('material_'):
        user_data[chat_id]['wall_material'] = call.data.split('_')[1]
        # Запрашиваем ширину стены после выбора материала
        msg = bot.send_message(chat_id, 'Введите ширину стены (м):')
        bot.register_next_step_handler(msg, wall_width_step)

    bot.answer_callback_query(call.id)  # Важно добавить для избежания "зависания" callback-уведомлений

def wall_width_step(message):
    chat_id = message.chat.id
    W = message.text
    user_data[chat_id]['wall_width'] = W
    # Вызываем функцию расчета
    city_name = user_data[chat_id]['city_name']
    wall_type = user_data[chat_id]['wall_type']
    wall_material = user_data[chat_id]['wall_material']
    variant = calculate_variant(wall_type, wall_material)
    thermal_technical_calculation(message, city_name, variant, W)
    # Очистка данных пользователя после завершения расчета
    del user_data[chat_id]



bot.polling(none_stop=True)